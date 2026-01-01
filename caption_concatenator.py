#!/usr/bin/env python3
"""
Caption Concatenator - Combine caption files into a single document ordered by date
"""

import json
import argparse
import re
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Optional, Tuple

# Optional imports for document generation
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from fpdf import FPDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False


def load_batch_file(json_path: Path) -> List[Dict]:
    """Load videos from a caption batch JSON file"""
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    return data.get('videos', [])


def parse_upload_date(date_str: Optional[str]) -> Optional[datetime]:
    """Parse upload date string to datetime"""
    if not date_str:
        return None

    # Try various formats with their expected string lengths
    formats = [
        ('%Y%m%d', 8),              # 20240115
        ('%Y-%m-%d', 10),           # 2024-01-15
        ('%Y-%m-%dT%H:%M:%S', 19),  # ISO format
    ]

    for fmt, length in formats:
        try:
            return datetime.strptime(date_str[:length], fmt)
        except (ValueError, TypeError):
            continue

    return None


def extract_date_from_title(title: str) -> Optional[datetime]:
    """
    Extract date from video title.

    Supports various formats commonly found in titles:
    - 2024-01-15, 2024/01/15, 2024.01.15
    - 01-15-2024, 01/15/2024, 01.15.2024
    - Jan 15, 2024 / January 15, 2024
    - 15 Jan 2024 / 15 January 2024
    """
    if not title:
        return None

    # Pattern 1: YYYY-MM-DD, YYYY/MM/DD, YYYY.MM.DD
    match = re.search(r'(20\d{2})[-/.](\d{1,2})[-/.](\d{1,2})', title)
    if match:
        try:
            year, month, day = int(match.group(1)), int(match.group(2)), int(match.group(3))
            return datetime(year, month, day)
        except ValueError:
            pass

    # Pattern 2: MM-DD-YYYY, MM/DD/YYYY, MM.DD.YYYY
    match = re.search(r'(\d{1,2})[-/.](\d{1,2})[-/.](20\d{2})', title)
    if match:
        try:
            month, day, year = int(match.group(1)), int(match.group(2)), int(match.group(3))
            return datetime(year, month, day)
        except ValueError:
            pass

    # Pattern 3: Month name DD, YYYY (e.g., "Jan 15, 2024" or "January 15, 2024")
    months = {
        'jan': 1, 'january': 1, 'feb': 2, 'february': 2, 'mar': 3, 'march': 3,
        'apr': 4, 'april': 4, 'may': 5, 'jun': 6, 'june': 6,
        'jul': 7, 'july': 7, 'aug': 8, 'august': 8, 'sep': 9, 'september': 9,
        'oct': 10, 'october': 10, 'nov': 11, 'november': 11, 'dec': 12, 'december': 12
    }

    match = re.search(
        r'(jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|june?|'
        r'july?|aug(?:ust)?|sep(?:tember)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?)'
        r'\s+(\d{1,2})(?:st|nd|rd|th)?,?\s*(20\d{2})',
        title, re.IGNORECASE
    )
    if match:
        try:
            month = months[match.group(1).lower()]
            day = int(match.group(2))
            year = int(match.group(3))
            return datetime(year, month, day)
        except (ValueError, KeyError):
            pass

    # Pattern 4: DD Month YYYY (e.g., "15 Jan 2024" or "15 January 2024")
    match = re.search(
        r'(\d{1,2})(?:st|nd|rd|th)?\s+'
        r'(jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|june?|'
        r'july?|aug(?:ust)?|sep(?:tember)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?)'
        r',?\s*(20\d{2})',
        title, re.IGNORECASE
    )
    if match:
        try:
            day = int(match.group(1))
            month = months[match.group(2).lower()]
            year = int(match.group(3))
            return datetime(year, month, day)
        except (ValueError, KeyError):
            pass

    return None


def get_video_date(video: Dict) -> Optional[datetime]:
    """
    Get the best available date for a video.
    Priority: upload_date > date extracted from title
    """
    # First try the upload_date field
    upload_date = parse_upload_date(video.get('upload_date'))
    if upload_date:
        return upload_date

    # Fall back to extracting from title
    title = video.get('title', '')
    return extract_date_from_title(title)


def format_date_str(date_str: Optional[str]) -> str:
    """Format date string for display"""
    dt = parse_upload_date(date_str)
    if dt:
        return dt.strftime('%Y-%m-%d')
    return 'Unknown date'


def format_video_date(video: Dict) -> str:
    """Format the best available date for a video"""
    dt = get_video_date(video)
    if dt:
        return dt.strftime('%Y-%m-%d')
    return 'Unknown date'


def fix_duplicate_words(text: str) -> str:
    """
    Find consecutive duplicate words and replace every other one with a comma.

    Examples:
        "the the quick" → "the, quick"
        "I I I think" → "I, I, think"
        "go go go go" → "go, go, go,"
    """
    if not text:
        return text

    words = text.split()
    if len(words) < 2:
        return text

    result = []
    i = 0

    while i < len(words):
        current_word = words[i]

        # Count consecutive duplicates
        j = i + 1
        while j < len(words) and words[j].lower() == current_word.lower():
            j += 1

        duplicate_count = j - i

        if duplicate_count == 1:
            # No duplicate, just add the word
            result.append(current_word)
        else:
            # Found duplicates - keep first, replace every other with comma
            for k in range(duplicate_count):
                if k % 2 == 0:
                    result.append(words[i + k])
                else:
                    result.append(',')

        i = j

    # Join and clean up spacing around commas
    output = ' '.join(result)
    # Fix spacing: "word , word" → "word, word"
    output = re.sub(r'\s+,', ',', output)
    # Remove trailing comma before period or end
    output = re.sub(r',(\s*[.!?])', r'\1', output)

    return output


def concatenate_text(
    json_files: List[Path],
    output_format: str = 'text',
    include_metadata: bool = True,
    reverse_order: bool = False,
    fix_duplicates: bool = False
) -> str:
    """
    Concatenate full text from all videos in JSON files.

    Args:
        json_files: List of caption batch JSON files
        output_format: 'text' or 'markdown'
        include_metadata: Include video title, URL, date
        reverse_order: If True, newest first; otherwise oldest first
        fix_duplicates: If True, replace duplicate consecutive words with commas

    Returns:
        Concatenated text
    """
    # Collect all videos
    all_videos = []
    for json_path in json_files:
        videos = load_batch_file(json_path)
        all_videos.extend(videos)

    # Sort by date (upload_date or extracted from title)
    def sort_key(video):
        dt = get_video_date(video)
        if dt:
            return dt
        # Put videos without dates at the end
        return datetime.max if not reverse_order else datetime.min

    all_videos.sort(key=sort_key, reverse=reverse_order)

    # Build output
    lines = []

    if output_format == 'markdown':
        lines.append('# Combined Captions')
        lines.append('')
        lines.append(f'*Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}*')
        lines.append(f'*Total videos: {len(all_videos)}*')
        lines.append('')

        for video in all_videos:
            title = video.get('title', 'Untitled')
            url = video.get('url', '')
            date = format_video_date(video)
            full_text = video.get('full_text', '')

            if fix_duplicates:
                full_text = fix_duplicate_words(full_text)

            if include_metadata:
                lines.append(f'## {title}')
                lines.append('')
                lines.append(f'**Date:** {date}  ')
                lines.append(f'**URL:** {url}')
                lines.append('')

            lines.append(full_text)
            lines.append('')
            lines.append('---')
            lines.append('')

    else:  # Plain text
        width = 80
        lines.append('=' * width)
        lines.append('COMBINED CAPTIONS')
        lines.append(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
        lines.append(f'Total videos: {len(all_videos)}')
        lines.append('=' * width)
        lines.append('')

        for video in all_videos:
            title = video.get('title', 'Untitled')
            url = video.get('url', '')
            date = format_video_date(video)
            full_text = video.get('full_text', '')

            if fix_duplicates:
                full_text = fix_duplicate_words(full_text)

            if include_metadata:
                lines.append('=' * width)
                lines.append(f'Video: {title}')
                lines.append(f'Date: {date}')
                lines.append(f'URL: {url}')
                lines.append('=' * width)
                lines.append('')

            lines.append(full_text)
            lines.append('')
            lines.append('')

    return '\n'.join(lines)


def prepare_videos(
    json_files: List[Path],
    reverse_order: bool = False,
    fix_duplicates: bool = True
) -> List[Dict]:
    """
    Load and prepare videos from JSON files.
    Returns sorted list of videos with processed text.
    """
    all_videos = []
    for json_path in json_files:
        videos = load_batch_file(json_path)
        all_videos.extend(videos)

    # Sort by date
    def sort_key(video):
        dt = get_video_date(video)
        if dt:
            return dt
        return datetime.max if not reverse_order else datetime.min

    all_videos.sort(key=sort_key, reverse=reverse_order)

    # Apply duplicate fixing if enabled
    if fix_duplicates:
        for video in all_videos:
            if 'full_text' in video:
                video['full_text'] = fix_duplicate_words(video['full_text'])

    return all_videos


def generate_toc(videos: List[Dict]) -> List[Tuple[str, str]]:
    """Generate table of contents entries: [(title, date), ...]"""
    toc = []
    for video in videos:
        title = video.get('title', 'Untitled')
        date = format_video_date(video)
        toc.append((title, date))
    return toc


def save_as_docx(
    videos: List[Dict],
    output_path: Path,
    title: str = "Combined Captions",
    include_toc: bool = True,
    include_metadata: bool = True
) -> bool:
    """Save videos as a Word document with table of contents."""
    if not DOCX_AVAILABLE:
        print("Error: python-docx not installed. Run: pip install python-docx")
        return False

    doc = Document()

    # Title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(f"Total videos: {len(videos)}")
    doc.add_paragraph()

    # Table of Contents
    if include_toc:
        doc.add_heading("Table of Contents", level=1)
        for i, video in enumerate(videos, 1):
            video_title = video.get('title', 'Untitled')
            date = format_video_date(video)
            toc_entry = doc.add_paragraph()
            toc_entry.add_run(f"{i}. {video_title}").bold = False
            toc_entry.add_run(f"  ({date})")
        doc.add_page_break()

    # Content
    for video in videos:
        video_title = video.get('title', 'Untitled')
        url = video.get('url', '')
        date = format_video_date(video)
        full_text = video.get('full_text', '')

        # Video heading
        doc.add_heading(video_title, level=1)

        if include_metadata:
            meta = doc.add_paragraph()
            meta.add_run("Date: ").bold = True
            meta.add_run(f"{date}\n")
            meta.add_run("URL: ").bold = True
            meta.add_run(url)

        # Content
        doc.add_paragraph(full_text)
        doc.add_paragraph()  # Spacing

    doc.save(str(output_path))
    return True


def save_as_pdf(
    videos: List[Dict],
    output_path: Path,
    title: str = "Combined Captions",
    include_toc: bool = True,
    include_metadata: bool = True
) -> bool:
    """Save videos as a PDF document with table of contents."""
    if not PDF_AVAILABLE:
        print("Error: fpdf2 not installed. Run: pip install fpdf2")
        return False

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)

    # Title page
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 24)
    pdf.cell(0, 40, title, ln=True, align="C")
    pdf.set_font("Helvetica", "", 12)
    pdf.cell(0, 10, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", ln=True, align="C")
    pdf.cell(0, 10, f"Total videos: {len(videos)}", ln=True, align="C")

    # Table of Contents
    if include_toc:
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 16)
        pdf.cell(0, 15, "Table of Contents", ln=True)
        pdf.set_font("Helvetica", "", 11)

        for i, video in enumerate(videos, 1):
            video_title = video.get('title', 'Untitled')[:60]  # Truncate long titles
            date = format_video_date(video)
            pdf.cell(0, 8, f"{i}. {video_title} ({date})", ln=True)

    # Content
    for video in videos:
        pdf.add_page()
        video_title = video.get('title', 'Untitled')
        url = video.get('url', '')
        date = format_video_date(video)
        full_text = video.get('full_text', '')

        # Video heading
        pdf.set_font("Helvetica", "B", 14)
        # Handle long titles by wrapping
        pdf.multi_cell(0, 8, video_title)

        if include_metadata:
            pdf.set_font("Helvetica", "", 10)
            pdf.cell(0, 6, f"Date: {date}", ln=True)
            pdf.set_text_color(0, 0, 255)
            pdf.cell(0, 6, url, ln=True)
            pdf.set_text_color(0, 0, 0)

        pdf.ln(5)

        # Content - handle encoding issues
        pdf.set_font("Helvetica", "", 11)
        # Replace problematic characters
        clean_text = full_text.encode('latin-1', errors='replace').decode('latin-1')
        pdf.multi_cell(0, 6, clean_text)

    pdf.output(str(output_path))
    return True


def main():
    parser = argparse.ArgumentParser(
        description='Concatenate caption files into a single document ordered by upload date',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Concatenate from a single batch file
  python caption_concatenator.py captions_batch_1.json -o combined.txt

  # Concatenate from multiple batch files
  python caption_concatenator.py batch1.json batch2.json batch3.json -o combined.txt

  # Output as markdown
  python caption_concatenator.py *.json -o combined.md --format markdown

  # Output both text and markdown
  python caption_concatenator.py *.json -o combined -f text,markdown

  # Output as PDF or Word document
  python caption_concatenator.py *.json -o combined.pdf -f pdf --title "My Captions"
  python caption_concatenator.py *.json -o combined.docx -f docx

  # Newest videos first
  python caption_concatenator.py *.json -o combined.txt --reverse

  # Text only, no metadata headers
  python caption_concatenator.py *.json -o combined.txt --no-metadata

  # Disable duplicate word fixing
  python caption_concatenator.py *.json -o combined.txt --no-dedup

  # Disable table of contents (PDF/Word only)
  python caption_concatenator.py *.json -o combined.pdf -f pdf --no-toc
        """
    )

    parser.add_argument('json_files', nargs='+', type=Path,
                       help='Caption batch JSON files to concatenate')
    parser.add_argument('-o', '--output', type=Path, required=True,
                       help='Output file path (extension added automatically for multiple formats)')
    parser.add_argument('-f', '--format', default='text',
                       help='Output format(s): text, markdown, pdf, docx (e.g., "text,markdown")')
    parser.add_argument('--title', default='Combined Captions',
                       help='Document title (for PDF/Word output)')
    parser.add_argument('--reverse', action='store_true',
                       help='Reverse order (newest first)')
    parser.add_argument('--no-metadata', action='store_true',
                       help='Exclude video titles, URLs, and dates')
    parser.add_argument('--no-dedup', action='store_true',
                       help='Disable fixing of consecutive duplicate words (enabled by default)')
    parser.add_argument('--no-toc', action='store_true',
                       help='Disable table of contents (PDF/Word only)')

    args = parser.parse_args()

    # Validate input files exist
    for json_file in args.json_files:
        if not json_file.exists():
            print(f"Error: File not found: {json_file}")
            return 1

    # Parse format(s)
    formats = [f.strip().lower() for f in args.format.split(',')]
    valid_formats = {'text', 'markdown', 'pdf', 'docx'}
    for fmt in formats:
        if fmt not in valid_formats:
            print(f"Error: Invalid format '{fmt}'. Use 'text', 'markdown', 'pdf', or 'docx'")
            return 1

    # Check for PDF/Word dependencies
    if 'pdf' in formats and not PDF_AVAILABLE:
        print("Error: PDF output requires fpdf2. Run: pip install fpdf2")
        return 1
    if 'docx' in formats and not DOCX_AVAILABLE:
        print("Error: Word output requires python-docx. Run: pip install python-docx")
        return 1

    # Auto-detect format from output extension if single format not specified
    if len(formats) == 1 and formats[0] == 'text':
        ext = args.output.suffix.lower()
        if ext == '.md':
            formats = ['markdown']
        elif ext == '.pdf':
            formats = ['pdf']
        elif ext == '.docx':
            formats = ['docx']

    print(f"Processing {len(args.json_files)} file(s)...")

    # Count videos
    total_videos = sum(len(load_batch_file(f)) for f in args.json_files)

    # Prepare videos once for all formats
    videos = prepare_videos(
        args.json_files,
        reverse_order=args.reverse,
        fix_duplicates=not args.no_dedup
    )

    # Generate output for each format
    output_files = []
    format_extensions = {'text': '.txt', 'markdown': '.md', 'pdf': '.pdf', 'docx': '.docx'}

    for output_format in formats:
        # Determine output path
        if len(formats) == 1:
            # Single format: use output path as-is
            output_path = args.output
        else:
            # Multiple formats: add extension to base name
            base = args.output.with_suffix('')  # Remove any existing extension
            output_path = base.with_suffix(format_extensions[output_format])

        # Generate output based on format
        if output_format == 'pdf':
            success = save_as_pdf(
                videos,
                output_path,
                title=args.title,
                include_toc=not args.no_toc,
                include_metadata=not args.no_metadata
            )
            if not success:
                return 1
        elif output_format == 'docx':
            success = save_as_docx(
                videos,
                output_path,
                title=args.title,
                include_toc=not args.no_toc,
                include_metadata=not args.no_metadata
            )
            if not success:
                return 1
        else:
            # Text or markdown
            result = concatenate_text(
                args.json_files,
                output_format=output_format,
                include_metadata=not args.no_metadata,
                reverse_order=args.reverse,
                fix_duplicates=not args.no_dedup
            )
            output_path.write_text(result, encoding='utf-8')

        output_files.append((output_format, output_path))

    print(f"Concatenated {total_videos} videos")
    for fmt, path in output_files:
        print(f"Output ({fmt}): {path}")
    if args.reverse:
        print("Order: Newest first")
    else:
        print("Order: Oldest first")

    return 0


if __name__ == '__main__':
    exit(main())
